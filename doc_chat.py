import io
from typing import List
import streamlit as st
from langchain.schema import Document
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
from langchain.chains import RetrievalQA
from langchain import LLMChain, PromptTemplate
from pypdf import PdfReader
import docx

EMBED_MODEL = "all-MiniLM-L6-v2"

# --------- Loaders ----------
def load_pdf_bytes(file_bytes: bytes, name="uploaded_pdf") -> List[Document]:
    reader = PdfReader(io.BytesIO(file_bytes))
    docs = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text()
        if txt and txt.strip():
            docs.append(Document(page_content=txt, metadata={"source": name, "page": i+1}))
    return docs

def load_docx_bytes(file_bytes: bytes, name="uploaded_docx") -> List[Document]:
    docs = []
    doc = docx.Document(io.BytesIO(file_bytes))
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t:
            docs.append(Document(page_content=t, metadata={"source": name, "para": i+1}))
    return docs

def load_txt_bytes(file_bytes: bytes, name="uploaded_txt") -> List[Document]:
    try:
        s = file_bytes.decode("utf-8")
    except:
        s = file_bytes.decode("latin-1")
    return [Document(page_content=s, metadata={"source": name})]

def load_documents_from_upload(uploaded_files) -> List[Document]:
    docs = []
    for f in uploaded_files:
        name = getattr(f, "name", "uploaded")
        data = f.read()
        if name.lower().endswith(".pdf"):
            docs.extend(load_pdf_bytes(data, name=name))
        elif name.lower().endswith(".docx") or name.lower().endswith(".doc"):
            docs.extend(load_docx_bytes(data, name=name))
        elif name.lower().endswith(".txt"):
            docs.extend(load_txt_bytes(data, name=name))
        else:
            # try pdf then txt fallback
            try:
                docs.extend(load_pdf_bytes(data, name=name))
            except Exception:
                docs.extend(load_txt_bytes(data, name=name))
    return docs

# --------- Embeddings & FAISS ----------
class EmbeddingManager:
    def __init__(self, model_name=EMBED_MODEL):
        self.model = SentenceTransformer(model_name)

    def embed_texts(self, texts: List[str]) -> np.ndarray:
        return self.model.encode(texts, show_progress_bar=False)

class FaissIndex:
    def __init__(self, vectors: np.ndarray, docs: List[Document]):
        d = vectors.shape[1]
        index = faiss.IndexFlatL2(d)
        index.add(vectors.astype("float32"))
        self.index = index
        self.docs = docs
        self.vectors = vectors

    @classmethod
    def from_texts(cls, texts: List[str], docs: List[Document], embed_mgr: EmbeddingManager):
        vecs = embed_mgr.embed_texts(texts)
        return cls(vecs, docs)

    def search(self, query_vec: np.ndarray, k=4):
        D, I = self.index.search(np.array([query_vec]).astype("float32"), k)
        hits = []
        for idx in I[0]:
            if idx < len(self.docs):
                hits.append(self.docs[idx])
        return hits

# --------- Retriever wrapper for LangChain ----------
class SimpleRetriever:
    def __init__(self, faiss_index: FaissIndex, embed_mgr: EmbeddingManager):
        self.faiss = faiss_index
        self.embed_mgr = embed_mgr

    def get_relevant_documents(self, query: str, k: int = 4):
        qv = self.embed_mgr.embed_texts([query])[0]
        return self.faiss.search(qv, k=k)

# --------- UI / Runner ----------
def run_doc_chat(llm):
    st.header("Document Q&A / Chat with Documents")
    st.write("Upload PDF / DOCX / TXT. The app will extract text, embed locally with sentence-transformers, index with FAISS, and answer questions using the LLM.")

    uploaded = st.file_uploader("Upload documents (multiple)", accept_multiple_files=True, key="docchat_upload")
    if not uploaded:
        st.info("Upload one or more documents to get started.")
        return

    with st.spinner("Loading documents..."):
        docs = load_documents_from_upload(uploaded)
    st.success(f"Loaded {len(docs)} text chunks from uploads.")

    # Prepare texts and embeddings
    texts = [d.page_content for d in docs]
    embed_mgr = EmbeddingManager()
    with st.spinner("Creating embeddings and FAISS index (local)..."):
        faiss_idx = FaissIndex.from_texts(texts, docs, embed_mgr)

    if "doc_chat_history" not in st.session_state:
        st.session_state.doc_chat_history = []

    question = st.text_input("Ask a question about uploaded documents", key="doc_q")
    if st.button("Ask") and question:
        retriever = SimpleRetriever(faiss_idx, embed_mgr)
        # Build RetrievalQA with LangChain (returns answer + sources)
        retriever_for_chain = retriever  # object shape matches get_relevant_documents
        chain = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever_for_chain, return_source_documents=True)
        result = chain(question)
        answer = result.get("result") or result.get("answer") or str(result)
        src_docs = result.get("source_documents", []) or result.get("source_documents", [])
        st.session_state.doc_chat_history.append({"q": question, "a": answer, "sources": src_docs})

    st.subheader("Chat History")
    for item in reversed(st.session_state.doc_chat_history[-20:]):
        st.markdown(f"**Q:** {item['q']}")
        st.markdown(f"**A:** {item['a']}")
        if item.get("sources"):
            st.write("**Sources:**")
            for s in item["sources"]:
                meta = getattr(s, "metadata", {}) or {}
                st.write(f"- {meta.get('source','unknown')} (meta: {meta})")
        st.markdown("---")
